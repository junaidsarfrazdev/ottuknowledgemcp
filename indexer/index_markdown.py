"""Standalone markdown / docx / xlsx indexer for internal docs.

Files are specified by absolute path in config.MARKDOWN_FILES.
All go into a single collection: config.INTERNAL_DOCS_COLLECTION.
Incremental: per-file SHA tracked in a sidecar metadata file at the project root.
"""
from __future__ import annotations

import hashlib
import json
from datetime import datetime, timezone
from pathlib import Path

import chromadb
from langchain_text_splitters import Language, RecursiveCharacterTextSplitter
from rich.console import Console

from . import config
from .embeddings import OllamaEmbeddings

console = Console()

_MD_SPLITTER = RecursiveCharacterTextSplitter.from_language(
    language=Language.MARKDOWN,
    chunk_size=config.CHUNK_SIZE,
    chunk_overlap=config.CHUNK_OVERLAP,
)
_DEFAULT_SPLITTER = RecursiveCharacterTextSplitter(
    chunk_size=config.CHUNK_SIZE, chunk_overlap=config.CHUNK_OVERLAP
)

_INTERNAL_META_PATH = config.PROJECT_ROOT / ".internal_docs_metadata.json"


def _file_sha(path: Path) -> str:
    h = hashlib.sha1()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(64 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _load_meta() -> dict:
    if not _INTERNAL_META_PATH.exists():
        return {}
    try:
        return json.loads(_INTERNAL_META_PATH.read_text())
    except json.JSONDecodeError:
        return {}


def _save_meta(data: dict) -> None:
    _INTERNAL_META_PATH.write_text(json.dumps(data, indent=2))


def _load_md(path: Path) -> list[tuple[str, dict]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [(p, {}) for p in _MD_SPLITTER.split_text(text) if p.strip()]


def _load_docx(path: Path) -> list[tuple[str, dict]]:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        console.print("[red]python-docx not installed. Install requirements.txt.[/red]")
        return []
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return [(p, {}) for p in _DEFAULT_SPLITTER.split_text(text) if p.strip()]


def _load_xlsx(path: Path) -> list[tuple[str, dict]]:
    try:
        from openpyxl import load_workbook  # type: ignore
    except ImportError:
        console.print("[red]openpyxl not installed. Install requirements.txt.[/red]")
        return []
    wb = load_workbook(filename=str(path), data_only=True, read_only=True)
    chunks: list[tuple[str, dict]] = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows_text: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                rows_text.append(",".join(cells))
        if not rows_text:
            continue
        sheet_text = f"# Sheet: {sheet}\n" + "\n".join(rows_text)
        for piece in _DEFAULT_SPLITTER.split_text(sheet_text):
            if piece.strip():
                chunks.append((piece, {"sheet_name": sheet}))
    return chunks


def _load(path: Path) -> list[tuple[str, dict]]:
    suf = path.suffix.lower()
    if suf == ".md":
        return _load_md(path)
    if suf == ".docx":
        return _load_docx(path)
    if suf == ".xlsx":
        return _load_xlsx(path)
    console.print(f"[yellow]⚠[/yellow]  Skipping unsupported file: {path}")
    return []


def index_internal_docs(
    client: chromadb.ClientAPI, embeddings: OllamaEmbeddings
) -> dict:
    if not config.MARKDOWN_FILES:
        console.print(
            "[dim]No MARKDOWN_FILES configured. Drop files in docs_local/ and add their "
            "absolute paths to indexer/config.py MARKDOWN_FILES to index them.[/dim]"
        )
        return {"chunks_written": 0}

    collection = client.get_or_create_collection(
        name=config.INTERNAL_DOCS_COLLECTION,
        metadata={"ottu_internal_docs": True},
    )

    meta_store = _load_meta()
    known_files: dict[str, dict] = meta_store.get("files", {})

    ids: list[str] = []
    docs: list[str] = []
    metas: list[dict] = []
    now_iso = datetime.now(timezone.utc).isoformat()

    current_paths: set[str] = set()
    changed = 0
    skipped = 0
    removed = 0

    for raw_path in config.MARKDOWN_FILES:
        path = Path(raw_path)
        if not path.exists():
            console.print(f"[red]✗[/red] missing internal doc: {path}")
            continue
        key = path.as_posix()
        current_paths.add(key)
        sha = _file_sha(path)
        prev = known_files.get(key)
        if prev and prev.get("sha") == sha:
            skipped += 1
            continue

        try:
            collection.delete(where={"path": str(path)})
        except Exception:
            pass

        pieces = _load(path)
        for idx, (piece, extra) in enumerate(pieces):
            cid = f"internal:{key}:{idx}"
            meta = {
                "path": str(path),
                "filename": path.name,
                "chunk_index": idx,
                "format": path.suffix.lower().lstrip("."),
                "file_sha": sha,
                "ingested_at": now_iso,
            }
            meta.update(extra)
            ids.append(cid)
            docs.append(piece)
            metas.append(meta)
        known_files[key] = {"sha": sha, "chunks": len(pieces)}
        changed += 1
        console.print(f"[green]✓[/green] {path.name}: {len(pieces)} chunks")

    # Remove chunks for files dropped from MARKDOWN_FILES or missing on disk
    stale = set(known_files.keys()) - current_paths
    for key in stale:
        try:
            collection.delete(where={"path": key})
        except Exception:
            pass
        known_files.pop(key, None)
        removed += 1

    # Upsert + per-batch checkpoint so partial failures don't poison
    # incremental state or collide on re-run.
    if ids:
        B = config.EMBED_BATCH_SIZE
        for i in range(0, len(ids), B):
            batch = docs[i : i + B]
            vectors = embeddings.embed_documents(batch, batch_size=B)
            collection.upsert(
                ids=ids[i : i + B],
                documents=batch,
                metadatas=metas[i : i + B],
                embeddings=vectors,
            )
            _save_meta({"indexed_at": now_iso, "files": known_files})

    _save_meta({"indexed_at": now_iso, "files": known_files})

    console.print(
        f"[green]✓[/green] internal docs: "
        f"{changed} changed, {skipped} unchanged, {removed} removed, "
        f"{len(ids)} chunks written"
    )
    return {
        "chunks_written": len(ids),
        "files_changed": changed,
        "files_skipped": skipped,
        "files_removed": removed,
    }
